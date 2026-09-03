"""多格式解析 → 结构化 Block 流。parse_file() 返回 (file_hash, [Segment])。

PDF 是重点，按页路由：
  - 数字原生页：PyMuPDF dict 提取 → 多栏恢复阅读顺序 → 表格/公式/代码识别
  - 扫描页（无文字层）：RapidOCR 渲染识别
  - 跨页表格按「表头重复」判据拼接后再输出
"""
from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

from mini_rag.config import settings
from mini_rag.core import cleaner

# 公式启发式：数学符号密度或 LaTeX 片段
_MATH_CHARS = set("∑∏∫∝∈∉∀∃∇≤≥≠±×÷√∞αβγδεθλμπσφψωΔΣΩΦΨ⊂⊃∪∩→←⇒⇔")
_MATH_RE = [
    re.compile(r"\$[^$]{1,120}\$"),
    re.compile(r"\\(?:begin|frac|sum|int|sqrt|alpha|beta|delta|sigma|omega)\b"),
    re.compile(r"[A-Za-z]\s*(?:\^\{?\d|\_\{?\d)"),
    re.compile(r"(?:^|\s)[A-Za-z]{1,3}\s*=\s*[^=]{1,40}(?:$|\s)"),
]
_CODE_FONT = re.compile(r"mono|courier|consolas|code|terminal", re.I)
# 目录页判定：显式标题 / 纯章节号 / 点线+页码
_TOC_HEAD = re.compile(r"^(?:table\s+of\s+)?contents?$", re.I)   # Table of contents / Contents
_TOC_NUM = re.compile(r"^\d+(?:\.\d+)+\.?$")                      # 1.1 / 6.2.3（必须带点，避免误伤图示标注的裸数字 1/2）
_TOC_DOT = re.compile(r"\.{4,}\s*\d+\s*$")                        # ......2


@dataclass
class Segment:
    page: int | None           # 起始物理页（1 起）；非 PDF 为 None
    page_end: int | None       # 结束物理页（跨页表格与起始不同）
    heading: str               # 章节路径，如 "Chapter 3 > 3.2 Revenue Analysis"
    text: str
    kind: str = "text"         # text | heading | table | formula | code
    level: int = 0             # 标题层级，0 表示正文


class ParseError(Exception):
    pass


def _file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(1 << 20), b""):
            h.update(block)
    return h.hexdigest()


def _read_text(path: Path) -> str:
    raw = path.read_bytes()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        from charset_normalizer import from_bytes
        best = from_bytes(raw).best()
        return str(best) if best else raw.decode("utf-8", errors="replace")


# ==================== PDF ====================
def _is_scan_page(page) -> bool:
    """无文字层且页面含图像 → 扫描页。

    必须同时要求「有图像」：真正的空白页（如分隔页）文字也少，但 OCR 它纯属浪费。
    """
    try:
        if len(page.get_text("text").strip()) >= settings.PDF_TEXT_LAYER_MIN_CHARS:
            return False
        return bool(page.get_images(full=False))
    except Exception:
        return False


def _is_toc_page(blocks: list[dict]) -> bool:
    """整页目录（TOC）判定，须在清洗前用原始块调用。

    目录页的正文块会带「点线 + 页码」，而 cleaner 会把点线行清掉，所以必须
    在 clean_blocks 之前判断，否则信号丢失、续页目录漏判。两个判据任一命中：
      1. 存在 "Table of Contents" / "Contents" 标题块；
      2. 目录条目行（纯章节号 / 点线+页码）占非空行的比例 ≥ TOC_LINE_RATIO。
    """
    if not settings.TOC_DETECT or not blocks:
        return False
    total = 0
    toc_lines = 0
    has_head = False
    for b in blocks:
        text = b["text"].strip()
        if not text:
            continue
        for ln in text.split("\n"):
            ln = ln.strip()
            if not ln:
                continue
            if _TOC_HEAD.fullmatch(ln):
                has_head = True
            total += 1
            if _TOC_NUM.match(ln) or _TOC_DOT.search(ln):
                toc_lines += 1
    if has_head:
        return True
    return total > 0 and toc_lines / total >= settings.TOC_LINE_RATIO


def _table_to_markdown(rows: list[list]) -> str:
    if not rows:
        return ""
    clean = [["" if c is None else str(c).replace("\n", " ").strip() for c in r]
             for r in rows]
    if not clean or not clean[0]:
        return ""
    head = clean[0]
    out = ["| " + " | ".join(head) + " |",
           "| " + " | ".join("---" for _ in head) + " |"]
    for r in clean[1:]:
        r = r + [""] * (len(head) - len(r)) if len(r) < len(head) else r[:len(head)]
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def _in_table(bbox, tables) -> bool:
    for t in tables:
        if bbox.intersects(t["bbox"]) and _overlap_ratio(bbox, t["bbox"]) > 0.5:
            return True
    return False


def _overlap_ratio(a, b) -> float:
    inter = a & b
    if inter.is_empty:
        return 0.0
    area = max(1.0, a.get_area())
    return inter.get_area() / area


def _reorder_columns(blocks: list[dict]) -> list[dict]:
    """贪心分栏恢复阅读顺序。

    按 (y0, x0) 排序后顺序扫描：若某块顶端高于「当前栏已到达的最低点」，
    说明视线跳回了页面上方 —— 即进入下一栏。最后按栏的左边界排序输出。
    """
    if (not settings.MULTICOLUMN_ENABLED
            or len(blocks) < settings.MULTICOLUMN_MIN_BLOCKS):
        return sorted(blocks, key=lambda b: (b["bbox"][1], b["bbox"][0]))
    ordered = sorted(blocks, key=lambda b: (b["bbox"][1], b["bbox"][0]))
    cols: list[list[dict]] = []
    cur = [ordered[0]]
    cur_bottom = ordered[0]["bbox"][3]
    for b in ordered[1:]:
        if b["bbox"][1] >= cur_bottom - 2.0:      # 仍在当前栏下方
            cur.append(b)
            cur_bottom = max(cur_bottom, b["bbox"][3])
        else:                                      # 回到上方 → 新栏
            cols.append(cur)
            cur = [b]
            cur_bottom = b["bbox"][3]
    cols.append(cur)
    cols.sort(key=lambda c: min(x["bbox"][0] for x in c))
    out: list[dict] = []
    for c in cols:
        out.extend(sorted(c, key=lambda b: b["bbox"][1]))
    return out


def _native_page(page) -> tuple[list[dict], list[dict]]:
    """提取单页文本块与表格。表格区域内的文本块会被排除，避免正文重复。"""
    import fitz
    tables: list[dict] = []
    if settings.TABLE_ENABLED:
        try:
            for tf in page.find_tables():
                rows = tf.extract()
                md = _table_to_markdown(rows)
                if not md:
                    continue
                tables.append({"rows": rows, "md": md,
                               "bbox": fitz.Rect(tf.bbox)})
        except Exception:
            pass                                   # 表格提取失败不影响正文

    d = page.get_text("dict")
    blocks: list[dict] = []
    for b in d.get("blocks", []):
        if b.get("type") != 0:
            continue
        lines = b.get("lines", [])
        # 按「行」拼接而非按 span：项目符号列表在 PDF 里是一行一个 bullet，
        # 直接拼 span 会把 •A•B•C 挤成一行，切分后完全不可读。
        text = "\n".join("".join(s["text"] for s in ln.get("spans", []))
                         for ln in lines)
        if not text.strip():
            continue
        spans = [s for ln in lines for s in ln.get("spans", [])]
        if not spans:
            continue
        size = max(s["size"] for s in spans)
        bold = any(s["flags"] & 16 for s in spans)
        fonts = {s.get("font", "") for s in spans}
        bbox = fitz.Rect(b["bbox"])
        if tables and _in_table(bbox, tables):
            continue
        blocks.append({
            "bbox": bbox, "text": text.strip(), "size": size, "bold": bold,
            "is_code": any(_CODE_FONT.search(f) for f in fonts),
        })
    return _reorder_columns(blocks), tables


_ocr_engine = None


def _get_ocr():
    global _ocr_engine
    if settings.OCR_PROVIDER == "llamaparse":
        return _llamaparse_engine()
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
    return _ocr_engine


def _llamaparse_engine():
    """LlamaParse 云端 OCR（预留桩，未实现）。

    LlamaParse 需出网调用 + API key，违背本地优先与隐私合规约束（DOC-05）。
    若要启用：`pip install llama-parse` 并配置 `LLAMA_CLOUD_API_KEY`，然后在此
    返回一个实现了 `(img: np.ndarray) -> [(box, text, conf), ...]` 的可调用
    对象（对齐 RapidOCR 的返回三元组），`_ocr_page` 即可无改动复用。
    """
    raise NotImplementedError(
        "LlamaParseOCR 为预留实现：需 `pip install llama-parse` 并配置 "
        "LLAMA_CLOUD_API_KEY；文档内容会出网，违反 DOC-05 本地隐私约束，"
        "请在明确接受后实现 `_llamaparse_engine()` 并返回可调用对象。")


def _ocr_page(page) -> list[dict]:
    """扫描页 OCR：渲染位图 → 识别 → 按 y 聚类成行 → 按 x 排序拼行。"""
    import numpy as np
    engine = _get_ocr()
    pix = page.get_pixmap(dpi=settings.OCR_DPI)
    img = np.frombuffer(pix.samples, dtype=np.uint8)
    img = img.reshape(pix.height, pix.width, pix.n)
    if pix.n == 4:
        img = img[:, :, :3]
    result, _ = engine(img)
    if not result:
        return []
    items = []
    for box, text, score in result:
        if not text or not str(text).strip():
            continue
        # RapidOCR 返回的置信度是字符串形式的浮点数，必须显式转换
        try:
            conf = float(score)
        except (TypeError, ValueError):
            conf = 1.0
        if conf < settings.OCR_MIN_CONFIDENCE:
            continue
        ys = [p[1] for p in box]
        xs = [p[0] for p in box]
        items.append((min(ys), min(xs), text.strip()))
    items.sort()
    rows: list[list[tuple[float, str]]] = []
    cur_y: float | None = None
    for y, x, t in items:
        if cur_y is None or abs(y - cur_y) <= 8:
            if not rows:
                rows.append([])
            rows[-1].append((x, t))
            cur_y = y if cur_y is None else cur_y
        else:
            rows.append([(x, t)])
            cur_y = y
    out = []
    for r in rows:
        line = " ".join(t for _, t in sorted(r)).strip()
        if line:
            out.append({"bbox": (0, 0, 0, 0), "text": line, "size": 0.0,
                        "bold": False, "is_code": False})
    return out


def _is_formula(text: str) -> bool:
    if not settings.FORMULA_DETECT or not text.strip():
        return False
    if any(r.search(text) for r in _MATH_RE):
        return True
    hits = sum(1 for c in text if c in _MATH_CHARS)
    return hits / max(1, len(text)) > 0.04


def _heading_thresholds(sizes: list[float]) -> list[float]:
    """从字号分布推断标题层级：正文取众数，大于正文 1.06 倍的字号降序排列。"""
    from collections import Counter
    if not sizes:
        return []
    cnt = Counter(round(s, 1) for s in sizes)
    body = cnt.most_common(1)[0][0]
    heads = sorted((s for s in cnt if s > body * 1.06), reverse=True)
    return heads[:6]


def _level_of(block: dict, thresholds: list[float]) -> int:
    if not thresholds or not block["size"]:
        return 0
    for i, t in enumerate(thresholds, 1):
        if block["size"] >= t:
            return i
    return 0


def _merge_cross_page(pages: list[dict]) -> None:
    """跨页表格拼接：下一页首个表格的表头与上一页末个表格表头一致 → 合并。"""
    if not settings.TABLE_CROSS_PAGE:
        return
    for i in range(len(pages) - 1):
        prev = pages[i]["tables"]
        nxt = pages[i + 1]["tables"]
        if not prev or not nxt:
            continue
        a, b = prev[-1], nxt[0]
        ra, rb = a["rows"], b["rows"]
        if not ra or not rb or len(ra[0]) != len(rb[0]):
            continue
        ka = [str(c or "").strip().lower() for c in ra[0]]
        kb = [str(c or "").strip().lower() for c in rb[0]]
        if not any(ka) or ka != kb:
            continue
        merged = ra + rb[1:]
        a["rows"] = merged
        a["md"] = _table_to_markdown(merged)
        a["page_end"] = pages[i + 1]["no"]
        nxt.pop(0)                                  # 续表已并入上一页


def _parse_pdf(path: Path) -> list[Segment]:
    import fitz
    doc = fitz.open(path)
    pages: list[dict] = []
    try:
        npages = doc.page_count
        if npages == 0:
            raise ParseError("空 PDF")
        scan = [_is_scan_page(p) for p in doc]
        n_scan = sum(scan)
        use_ocr = (settings.OCR_ENABLED and n_scan > 0
                   and n_scan <= settings.OCR_MAX_PAGES)
        if n_scan and not use_ocr:
            print(f"    [warn] {n_scan} 个扫描页超过 OCR_MAX_PAGES="
                  f"{settings.OCR_MAX_PAGES}，已跳过 OCR")

        sizes: list[float] = []
        for i, page in enumerate(doc):
            if scan[i] and use_ocr:
                blocks = _ocr_page(page)
                tables: list[dict] = []
            else:
                blocks, tables = _native_page(page)
                sizes.extend(b["size"] for b in blocks)
            for t in tables:
                t.setdefault("page_start", i + 1)
                t.setdefault("page_end", i + 1)
            pages.append({"no": i + 1, "blocks": blocks, "tables": tables,
                          "height": page.rect.height, "scan": scan[i],
                          "is_toc": _is_toc_page(blocks)})

        _merge_cross_page(pages)
        thresholds = _heading_thresholds(sizes)
        noise = cleaner.learn_noise(pages)
        for pg in pages:
            pg["blocks"] = cleaner.clean_blocks(pg["blocks"], noise, pg["height"])
    finally:
        doc.close()

    segs: list[Segment] = []
    stack: list[tuple[int, str]] = []
    for pg in pages:
        if pg.get("is_toc"):                    # 目录页整页跳过：导航不是正文
            continue
        path_str = " > ".join(t for _, t in stack)
        for b in pg["blocks"]:
            text = b["text"]
            if not text:
                continue
            lvl = _level_of(b, thresholds)
            if lvl:
                title = " ".join(text.split("\n")[0].split())[:120]
                if title:
                    stack = [(l, t) for l, t in stack if l < lvl]
                    stack.append((lvl, title))
                    path_str = " > ".join(t for _, t in stack)
                    segs.append(Segment(pg["no"], pg["no"], path_str, title,
                                        "heading", lvl))
                continue
            kind = "code" if b.get("is_code") else (
                "formula" if _is_formula(text) else "text")
            segs.append(Segment(pg["no"], pg["no"], path_str, text, kind,
                                len(stack)))
        for t in pg["tables"]:
            segs.append(Segment(t["page_start"], t["page_end"], path_str,
                                t["md"], "table", len(stack)))
    if not segs:
        raise ParseError("无有效内容（可能全为扫描页且未启用 OCR）")
    return segs


# ==================== 其他格式 ====================
def _md_segments(text: str) -> list[Segment]:
    """Markdown → Segment。围栏代码块单独成 code 段，保证切分时不与正文混切。"""
    segs: list[Segment] = []
    stack: list[tuple[int, str]] = []
    buf: list[str] = []
    code: list[str] = []
    in_code = False
    heading = ""

    def flush_text() -> None:
        nonlocal buf
        body = cleaner.clean_text("\n".join(buf)).strip()
        if body:
            segs.append(Segment(None, None, heading, body, "text", len(stack)))
        buf = []

    def flush_code() -> None:
        nonlocal code
        body = "\n".join(code).strip()
        if body:
            segs.append(Segment(None, None, heading, body, "code", len(stack)))
        code = []

    for line in text.splitlines():
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
            else:
                flush_text()
            in_code = not in_code
            continue
        if in_code:
            code.append(line)
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            flush_text()
            level = len(m.group(1))
            title = m.group(2).strip()
            stack = [(l, t) for l, t in stack if l < level]
            stack.append((level, title))
            heading = " > ".join(t for _, t in stack)
            segs.append(Segment(None, None, heading, title, "heading", level))
        else:
            buf.append(line)
    if in_code:
        flush_code()
    flush_text()
    return segs


def _parse_md(path: Path) -> list[Segment]:
    return _md_segments(_read_text(path))


def _parse_txt(path: Path) -> list[Segment]:
    text = cleaner.clean_text(_read_text(path))
    return [Segment(None, None, "", text)] if text.strip() else []


def _parse_docx(path: Path) -> list[Segment]:
    import docx
    d = docx.Document(str(path))
    parts: list[str] = []
    for para in d.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    tables: list[Segment] = []
    for table in d.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        md = _table_to_markdown(rows)
        if md:
            tables.append(Segment(None, None, "", md, "table"))
    text = cleaner.clean_text("\n".join(parts))
    segs = [Segment(None, None, "", text)] if text.strip() else []
    return segs + tables


def _parse_html(path: Path) -> list[Segment]:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(_read_text(path), "lxml")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    root = soup.find("article") or soup.find("main") or soup.find("body") or soup
    text = cleaner.clean_text(root.get_text("\n", strip=True))
    return [Segment(None, None, "", text)] if text.strip() else []


_PARSERS = {
    ".pdf": _parse_pdf,
    ".md": _parse_md,
    ".txt": _parse_txt,
    ".docx": _parse_docx,
    ".html": _parse_html,
    ".htm": _parse_html,
}


def file_hash(path: str | Path) -> str:
    """公开入口：pipeline 幂等检查要先算哈希，避免重复读文件。"""
    return _file_hash(Path(path))


def pdf_page_count(path: str | Path) -> int:
    """取 PDF 页数用于切分档位选择。非 PDF 或读取失败返回 0。"""
    p = Path(path)
    if p.suffix.lower() != ".pdf":
        return 0
    try:
        import fitz
        with fitz.open(p) as d:
            return d.page_count
    except Exception:
        return 0


def parse_file(path: str | Path,
               file_hash: str | None = None) -> tuple[str, list[Segment]]:
    """解析文件 → (file_hash, segments)。失败抛 ParseError，由 pipeline 捕获记录。

    file_hash 可由调用方传入以避免重复读文件算哈希（pipeline 幂等检查已算过一次）。
    """
    p = Path(path)
    fn = _PARSERS.get(p.suffix.lower())
    if fn is None:
        raise ParseError(f"unsupported_format: {p.suffix.lower()}")
    return file_hash or _file_hash(p), fn(p)
